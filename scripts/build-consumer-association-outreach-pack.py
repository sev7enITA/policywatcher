from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "output" / "outreach"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_PATH = OUTPUT_DIR / "PolicyWatcher-proposta-associazioni-consumatori-IT.docx"
WORDMARK = ROOT / "public" / "press-kit" / "policywatcher-wordmark-dark-2400x600.png"

BLUE = "2E74B5"
DEEP_BLUE = "12304A"
PALE_BLUE = "EAF3F9"
PALE_GRAY = "F2F4F7"
MID_GRAY = "586774"
LIGHT_RULE = "CBD5DE"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=LIGHT_RULE, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def set_no_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def style_run(run, size=11, bold=False, color=None, italic=False):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_body(doc, text, *, bold_prefix=None, after=6, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if bold_prefix and text.startswith(bold_prefix):
        head = p.add_run(bold_prefix)
        style_run(head, bold=True)
        tail = p.add_run(text[len(bold_prefix):])
        style_run(tail)
    else:
        run = p.add_run(text)
        style_run(run)
    return p


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    style_run(run, size=13, bold=True, color=BLUE)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.50)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    run = p.add_run(text)
    style_run(run)
    return p


def add_cell_text(cell, label, text, *, fill=WHITE):
    set_cell_shading(cell, fill)
    set_cell_margins(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    label_run = p.add_run(label.upper() + "\n")
    style_run(label_run, size=8.5, bold=True, color=BLUE)
    text_run = p.add_run(text)
    style_run(text_run, size=10.5, color=DEEP_BLUE)


def add_stat_cell(cell, value, label):
    set_cell_shading(cell, PALE_GRAY)
    set_cell_margins(cell, top=90, bottom=90)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    number = p.add_run(value + "\n")
    style_run(number, size=16, bold=True, color=DEEP_BLUE)
    caption = p.add_run(label)
    style_run(caption, size=8.5, color=MID_GRAY)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.72)
section.bottom_margin = Inches(0.68)
section.left_margin = Inches(0.82)
section.right_margin = Inches(0.82)
section.header_distance = Inches(0.30)
section.footer_distance = Inches(0.30)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

for style_name, size, color, before, after in (
    ("Title", 24, "000000", 0, 4),
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, "1F4D78", 8, 4),
):
    style = styles[style_name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

header = section.header
header.is_linked_to_previous = False
hp = header.paragraphs[0]
hp.paragraph_format.space_after = Pt(0)
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
hr = hp.add_run("POLICYWATCHER  |  PROPOSTA DI COLLABORAZIONE CIVICA")
style_run(hr, size=8, bold=True, color=MID_GRAY)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.paragraph_format.space_before = Pt(0)
fp.paragraph_format.space_after = Pt(0)
fr = fp.add_run("policywatcher.online  •  info@policywatcher.online  •  Pagina ")
style_run(fr, size=8, color=MID_GRAY)
add_page_field(fp)

logo_p = doc.add_paragraph()
logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
logo_p.paragraph_format.space_after = Pt(8)
logo_p.add_run().add_picture(str(WORDMARK), width=Inches(2.55))

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(3)
run = title.add_run("Un osservatorio civico sui contratti digitali")
style_run(run, size=22, bold=True, color="000000")

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(12)
run = subtitle.add_run("Proposta di collaborazione per le associazioni italiane dei consumatori")
style_run(run, size=13, color=MID_GRAY)

meta = doc.add_table(rows=2, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
meta.autofit = False
meta.columns[0].width = Inches(3.38)
meta.columns[1].width = Inches(3.38)
set_table_borders(meta, color=LIGHT_RULE, size="3")
add_cell_text(meta.cell(0, 0), "Proponente", "Fabrizio Degni · creatore e maintainer")
add_cell_text(meta.cell(0, 1), "Destinatari", "Presidenza, progetti e politiche dei consumatori")
add_cell_text(meta.cell(1, 0), "Contatto", "info@policywatcher.online")
add_cell_text(meta.cell(1, 1), "Data", "5 agosto 2026")

lead = doc.add_table(rows=1, cols=1)
lead.alignment = WD_TABLE_ALIGNMENT.CENTER
lead.autofit = False
lead.columns[0].width = Inches(6.76)
set_table_borders(lead, color=PALE_BLUE, size="0")
cell = lead.cell(0, 0)
set_cell_shading(cell, PALE_BLUE)
set_cell_margins(cell, top=130, bottom=130, start=180, end=180)
p = cell.paragraphs[0]
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.10
r = p.add_run(
    "PolicyWatcher è un progetto civic-tech indipendente che rende più leggibili, tracciabili e verificabili "
    "i cambiamenti nelle policy pubbliche dei servizi digitali. L’obiettivo è aiutare cittadini e organizzazioni "
    "a distinguere ciò che è documentato da ciò che non è ancora disponibile o verificato."
)
style_run(r, size=11, bold=True, color=DEEP_BLUE)

add_heading(doc, "La proposta")
proposal = doc.add_table(rows=1, cols=3)
proposal.alignment = WD_TABLE_ALIGNMENT.CENTER
proposal.autofit = False
for idx, width in enumerate((2.25, 2.25, 2.26)):
    proposal.columns[idx].width = Inches(width)
set_table_borders(proposal, color=LIGHT_RULE, size="3")
add_cell_text(proposal.cell(0, 0), "1 · Osservare", "Una watchlist condivisa di servizi digitali rilevanti per gli associati.", fill=PALE_GRAY)
add_cell_text(proposal.cell(0, 1), "2 · Verificare", "Fonti, timestamp e stati mancanti restano espliciti prima della pubblicazione.", fill=PALE_GRAY)
add_cell_text(proposal.cell(0, 2), "3 · Informare", "Digest e schede comprensibili, sempre con revisione umana dell’associazione.", fill=PALE_GRAY)

add_heading(doc, "Pilot di 60 giorni, senza impegni")
add_body(
    doc,
    "Co-progettare un test su 10–20 servizi scelti insieme: definizione del perimetro, attivazione del monitoraggio, "
    "due momenti di revisione e un report finale con evidenze, limiti e possibili sviluppi. Nessun contenuto verrebbe "
    "presentato come valutazione legale o pubblicato a nome dell’associazione senza approvazione."
)

stats = doc.add_table(rows=1, cols=4)
stats.alignment = WD_TABLE_ALIGNMENT.CENTER
stats.autofit = False
for idx in range(4):
    stats.columns[idx].width = Inches(1.69)
set_table_borders(stats, color=WHITE, size="6")
add_stat_cell(stats.cell(0, 0), "16", "aziende configurate")
add_stat_cell(stats.cell(0, 1), "6", "settori configurati")
add_stat_cell(stats.cell(0, 2), "15", "KPI canonici")
add_stat_cell(stats.cell(0, 3), "EN / IT", "lingue editoriali")

note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
note.paragraph_format.space_before = Pt(2)
note.paragraph_format.space_after = Pt(8)
nr = note.add_run(
    "Dati di perimetro al 2 agosto 2026: descrivono l’inventario configurato, non una copertura esaustiva né un punteggio di conformità."
)
style_run(nr, size=8.5, italic=True, color=MID_GRAY)

cta = doc.add_table(rows=1, cols=1)
cta.alignment = WD_TABLE_ALIGNMENT.CENTER
cta.autofit = False
cta.columns[0].width = Inches(6.76)
set_table_borders(cta, color=DEEP_BLUE, size="4")
cell = cta.cell(0, 0)
set_cell_shading(cell, DEEP_BLUE)
set_cell_margins(cell, top=130, bottom=130, start=180, end=180)
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(0)
r = p.add_run("Prossimo passo: un confronto esplorativo di 30 minuti per scegliere un caso d’uso concreto.")
style_run(r, size=11, bold=True, color=WHITE)

boundary = doc.add_paragraph()
boundary.alignment = WD_ALIGN_PARAGRAPH.CENTER
boundary.paragraph_format.space_before = Pt(7)
boundary.paragraph_format.space_after = Pt(0)
br = boundary.add_run(
    "PolicyWatcher è uno strumento informativo e di ricerca. Non fornisce consulenza legale, non certifica la conformità "
    "e non rivendica una copertura completa del mercato. Codice e materiali pubblici: github.com/sev7enITA/policywatcher"
)
style_run(br, size=8.5, color=MID_GRAY)

doc.save(OUTPUT_PATH)
print(OUTPUT_PATH)
