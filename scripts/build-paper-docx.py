#!/usr/bin/env python3
"""Apply PolicyWatcher publication styling to a Pandoc-generated DOCX."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEAL = "138A8A"
TEXT = "24364B"
MUTED = "5F6F82"
PALE = "F4F7FA"
RULE = "CCD7E2"


def set_cell_width(cell, twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(twips))
    tc_w.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=110, bottom=100, end=110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), RULE)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)


def style_run(run, size: float, color: str = TEXT, bold: bool | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def add_cover(doc: Document, logo_path: Path) -> None:
    title = doc.paragraphs[0]
    logo_p = title.insert_paragraph_before()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_p.paragraph_format.space_after = Pt(20)
    logo_run = logo_p.add_run()
    picture = logo_run.add_picture(str(logo_path), width=Inches(1.85))
    picture._inline.docPr.set("descr", "PolicyWatcher logo")

    kicker = title.insert_paragraph_before("SYSTEMS PAPER  |  RELEASE 3.9.0 BETA 21")
    kicker.alignment = WD_ALIGN_PARAGRAPH.LEFT
    kicker.paragraph_format.space_after = Pt(12)
    for run in kicker.runs:
        style_run(run, 9.5, TEAL, True)
        run.font.all_caps = True

    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(22)
    for run in title.runs:
        style_run(run, 27, DARK_BLUE, True)

    if len(doc.paragraphs) > 3:
        author = doc.paragraphs[3]
        author.alignment = WD_ALIGN_PARAGRAPH.LEFT
        author.paragraph_format.space_after = Pt(10)
        for run in author.runs:
            style_run(run, 11.5, TEXT)

    if len(doc.paragraphs) > 4:
        date = doc.paragraphs[4]
        date.alignment = WD_ALIGN_PARAGRAPH.LEFT
        date.paragraph_format.space_before = Pt(4)
        date.paragraph_format.space_after = Pt(0)
        for run in date.runs:
            style_run(run, 10.5, MUTED)
        date.add_run().add_break(WD_BREAK.PAGE)


def style_document(input_path: Path, output_path: Path, logo_path: Path) -> None:
    doc = Document(str(input_path))
    section = doc.sections[0]
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.92)
    section.right_margin = Inches(0.92)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.widow_control = True

    for style_name in ("Body Text", "First Paragraph", "Abstract"):
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = "Calibri"
            style.font.size = Pt(10.5)
            style.font.color.rgb = RGBColor.from_string(TEXT)
            style.paragraph_format.space_after = Pt(7)
            style.paragraph_format.line_spacing = 1.18
            style.paragraph_format.widow_control = True

    heading_tokens = {
        "Heading 1": (17, BLUE, 18, 7),
        "Heading 2": (13.5, DARK_BLUE, 13, 5),
        "Heading 3": (11.5, TEAL, 10, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        if name not in doc.styles:
            continue
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    # Pandoc maps LaTeX \paragraph headings to Heading 4. In this manuscript
    # they are the first subsection level below a section, so promote them to
    # Heading 2 to preserve a sequential accessibility hierarchy.
    for paragraph in doc.paragraphs:
        if paragraph.style.name == "Heading 4":
            paragraph.style = doc.styles["Heading 2"]

    if "Abstract Title" in doc.styles:
        style = doc.styles["Abstract Title"]
        style.font.name = "Calibri"
        style.font.size = Pt(14)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(BLUE)
        style.paragraph_format.space_after = Pt(7)
        style.paragraph_format.keep_with_next = True

    if "Title" in doc.styles:
        style = doc.styles["Title"]
        style.font.name = "Calibri"
        style.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    if "Author" in doc.styles:
        doc.styles["Author"].font.name = "Calibri"
    if "Date" in doc.styles:
        doc.styles["Date"].font.name = "Calibri"

    add_cover(doc, logo_path)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.text = "POLICYWATCHER  |  SYSTEMS PAPER  |  BETA 21"
    for run in header.runs:
        style_run(run, 8.5, MUTED, True)
    header.paragraph_format.space_after = Pt(2)

    first_footer = section.first_page_footer.paragraphs[0]
    first_footer.text = "policywatcher.online  |  30 July 2026"
    first_footer.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in first_footer.runs:
        style_run(run, 8.5, MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = "PolicyWatcher  |  Evidence-preserving policy monitoring     "
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in footer.runs:
        style_run(run, 8.2, MUTED)
    add_page_field(footer)

    printable_twips = 9590
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        set_table_borders(table)
        columns = len(table.columns)
        if columns == 2:
            widths = [int(printable_twips * 0.28), int(printable_twips * 0.72)]
        elif columns == 3:
            widths = [int(printable_twips * 0.22), int(printable_twips * 0.34), int(printable_twips * 0.44)]
        elif columns == 4:
            widths = [int(printable_twips * 0.19), int(printable_twips * 0.14), int(printable_twips * 0.14), int(printable_twips * 0.53)]
        else:
            widths = [int(printable_twips / columns)] * columns

        for row_index, row in enumerate(table.rows):
            prevent_row_split(row)
            if row_index == 0:
                set_repeat_table_header(row)
            for column_index, cell in enumerate(row.cells):
                set_cell_width(cell, widths[min(column_index, len(widths) - 1)])
                set_cell_margins(cell)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                if row_index == 0:
                    shade_cell(cell, BLUE)
                elif row_index % 2 == 0:
                    shade_cell(cell, PALE)
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(3)
                    paragraph.paragraph_format.line_spacing = 1.05
                    paragraph.paragraph_format.widow_control = True
                    for run in paragraph.runs:
                        style_run(run, 8.6, "FFFFFF" if row_index == 0 else TEXT, row_index == 0)

    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.widow_control = True
        if paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True

    doc.core_properties.title = "PolicyWatcher: An Evidence-Preserving Pipeline for Monitoring Terms-of-Service and Privacy-Policy Changes"
    doc.core_properties.subject = "PolicyWatcher systems paper, release 3.9.0 Beta 21"
    doc.core_properties.author = "Fabrizio Degni"
    doc.core_properties.keywords = "PolicyWatcher, policy monitoring, data quality, explainability, provenance, source reliability"
    doc.core_properties.comments = "Publication edition generated from policywatcher-arxiv.tex."
    doc.save(str(output_path))


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("--logo", required=True, type=Path)
    args = parser.parse_args()
    style_document(args.input, args.output, args.logo)


if __name__ == "__main__":
    main()
